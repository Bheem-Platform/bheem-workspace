"""
Bheem Docs - Export Service
============================
Exports documents to various formats: PDF, DOCX, HTML, Markdown.
Uses WeasyPrint for PDF and python-docx for DOCX generation.

Database Tables:
- workspace.docs_exports - Export history
"""

from typing import Optional, List, Dict, Any
from uuid import UUID
from datetime import datetime
import logging
import tempfile
import os
import json
import io
import re
import psycopg2
from psycopg2.extras import RealDictCursor

from core.config import settings

logger = logging.getLogger(__name__)


class DocsExportService:
    """
    Document export service.

    Converts Tiptap JSON content to:
    - PDF (via WeasyPrint)
    - DOCX (via python-docx)
    - HTML
    - Markdown
    """

    def __init__(self):
        """Initialize with database connection."""
        self.db_config = {
            'host': settings.ERP_DB_HOST,
            'port': settings.ERP_DB_PORT,
            'database': settings.ERP_DB_NAME,
            'user': settings.ERP_DB_USER,
            'password': settings.ERP_DB_PASSWORD,
        }

    def _get_connection(self):
        """Get database connection."""
        return psycopg2.connect(**self.db_config)

    # =========================================================================
    # TIPTAP TO HTML CONVERSION
    # =========================================================================

    def tiptap_to_html(self, content: Dict[str, Any]) -> str:
        """
        Convert Tiptap JSON to HTML.

        Args:
            content: Tiptap document JSON

        Returns:
            HTML string
        """
        if not content or 'content' not in content:
            return ''

        html_parts = []
        for node in content.get('content', []):
            html_parts.append(self._node_to_html(node))

        return '\n'.join(html_parts)

    def _node_to_html(self, node: Dict[str, Any]) -> str:
        """Convert a single Tiptap node to HTML."""
        node_type = node.get('type', '')
        attrs = node.get('attrs', {})
        content = node.get('content', [])

        if node_type == 'doc':
            return '\n'.join(self._node_to_html(c) for c in content)

        elif node_type == 'paragraph':
            inner = self._render_inline_content(content)
            return f'<p>{inner}</p>'

        elif node_type == 'heading':
            level = attrs.get('level', 1)
            inner = self._render_inline_content(content)
            return f'<h{level}>{inner}</h{level}>'

        elif node_type == 'bulletList':
            items = '\n'.join(self._node_to_html(c) for c in content)
            return f'<ul>\n{items}\n</ul>'

        elif node_type == 'orderedList':
            items = '\n'.join(self._node_to_html(c) for c in content)
            start = attrs.get('start', 1)
            return f'<ol start="{start}">\n{items}\n</ol>'

        elif node_type == 'listItem':
            inner = '\n'.join(self._node_to_html(c) for c in content)
            return f'<li>{inner}</li>'

        elif node_type == 'taskList':
            items = '\n'.join(self._node_to_html(c) for c in content)
            return f'<ul class="task-list">\n{items}\n</ul>'

        elif node_type == 'taskItem':
            checked = attrs.get('checked', False)
            checkbox = '<input type="checkbox" disabled' + (' checked' if checked else '') + '>'
            inner = '\n'.join(self._node_to_html(c) for c in content)
            return f'<li class="task-item">{checkbox} {inner}</li>'

        elif node_type == 'blockquote':
            inner = '\n'.join(self._node_to_html(c) for c in content)
            return f'<blockquote>{inner}</blockquote>'

        elif node_type == 'codeBlock':
            language = attrs.get('language', '')
            code = self._render_inline_content(content)
            return f'<pre><code class="language-{language}">{self._escape_html(code)}</code></pre>'

        elif node_type == 'horizontalRule':
            return '<hr>'

        elif node_type == 'image':
            src = attrs.get('src', '')
            alt = attrs.get('alt', '')
            title = attrs.get('title', '')
            return f'<img src="{src}" alt="{alt}" title="{title}">'

        elif node_type == 'table':
            rows = '\n'.join(self._node_to_html(c) for c in content)
            return f'<table>\n{rows}\n</table>'

        elif node_type == 'tableRow':
            cells = ''.join(self._node_to_html(c) for c in content)
            return f'<tr>{cells}</tr>'

        elif node_type == 'tableHeader':
            inner = '\n'.join(self._node_to_html(c) for c in content)
            colspan = attrs.get('colspan', 1)
            rowspan = attrs.get('rowspan', 1)
            attrs_str = f' colspan="{colspan}"' if colspan > 1 else ''
            attrs_str += f' rowspan="{rowspan}"' if rowspan > 1 else ''
            return f'<th{attrs_str}>{inner}</th>'

        elif node_type == 'tableCell':
            inner = '\n'.join(self._node_to_html(c) for c in content)
            colspan = attrs.get('colspan', 1)
            rowspan = attrs.get('rowspan', 1)
            attrs_str = f' colspan="{colspan}"' if colspan > 1 else ''
            attrs_str += f' rowspan="{rowspan}"' if rowspan > 1 else ''
            return f'<td{attrs_str}>{inner}</td>'

        elif node_type == 'text':
            text = node.get('text', '')
            text = self._escape_html(text)
            marks = node.get('marks', [])
            for mark in marks:
                text = self._apply_mark(text, mark)
            return text

        elif node_type == 'hardBreak':
            return '<br>'

        elif node_type == 'mention':
            label = attrs.get('label', attrs.get('id', ''))
            return f'<span class="mention">@{label}</span>'

        else:
            # Unknown node type, try to render content
            if content:
                return '\n'.join(self._node_to_html(c) for c in content)
            return ''

    def _render_inline_content(self, content: List[Dict]) -> str:
        """Render inline content (text with marks)."""
        parts = []
        for node in content:
            if node.get('type') == 'text':
                text = node.get('text', '')
                text = self._escape_html(text)
                for mark in node.get('marks', []):
                    text = self._apply_mark(text, mark)
                parts.append(text)
            else:
                parts.append(self._node_to_html(node))
        return ''.join(parts)

    def _apply_mark(self, text: str, mark: Dict) -> str:
        """Apply a mark to text."""
        mark_type = mark.get('type', '')
        attrs = mark.get('attrs', {})

        if mark_type == 'bold':
            return f'<strong>{text}</strong>'
        elif mark_type == 'italic':
            return f'<em>{text}</em>'
        elif mark_type == 'underline':
            return f'<u>{text}</u>'
        elif mark_type == 'strike':
            return f'<s>{text}</s>'
        elif mark_type == 'code':
            return f'<code>{text}</code>'
        elif mark_type == 'link':
            href = attrs.get('href', '')
            target = attrs.get('target', '_blank')
            return f'<a href="{href}" target="{target}">{text}</a>'
        elif mark_type == 'highlight':
            color = attrs.get('color', 'yellow')
            return f'<mark style="background-color: {color}">{text}</mark>'
        elif mark_type == 'textStyle':
            styles = []
            if 'color' in attrs:
                styles.append(f"color: {attrs['color']}")
            if 'fontSize' in attrs:
                styles.append(f"font-size: {attrs['fontSize']}")
            style_str = '; '.join(styles)
            return f'<span style="{style_str}">{text}</span>'
        else:
            return text

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;'))

    # =========================================================================
    # TIPTAP TO MARKDOWN CONVERSION
    # =========================================================================

    def tiptap_to_markdown(self, content: Dict[str, Any]) -> str:
        """
        Convert Tiptap JSON to Markdown.

        Args:
            content: Tiptap document JSON

        Returns:
            Markdown string
        """
        if not content or 'content' not in content:
            return ''

        md_parts = []
        for node in content.get('content', []):
            md_parts.append(self._node_to_markdown(node))

        return '\n\n'.join(filter(None, md_parts))

    def _node_to_markdown(self, node: Dict[str, Any], list_level: int = 0) -> str:
        """Convert a single Tiptap node to Markdown."""
        node_type = node.get('type', '')
        attrs = node.get('attrs', {})
        content = node.get('content', [])

        if node_type == 'doc':
            return '\n\n'.join(self._node_to_markdown(c) for c in content)

        elif node_type == 'paragraph':
            return self._render_inline_markdown(content)

        elif node_type == 'heading':
            level = attrs.get('level', 1)
            inner = self._render_inline_markdown(content)
            return f"{'#' * level} {inner}"

        elif node_type == 'bulletList':
            items = []
            for c in content:
                item_md = self._node_to_markdown(c, list_level)
                items.append(item_md)
            return '\n'.join(items)

        elif node_type == 'orderedList':
            items = []
            for i, c in enumerate(content, 1):
                item_md = self._node_to_markdown(c, list_level)
                items.append(item_md.replace('- ', f'{i}. ', 1))
            return '\n'.join(items)

        elif node_type == 'listItem':
            indent = '  ' * list_level
            inner_parts = []
            for c in content:
                if c.get('type') in ('bulletList', 'orderedList'):
                    inner_parts.append('\n' + self._node_to_markdown(c, list_level + 1))
                else:
                    inner_parts.append(self._node_to_markdown(c, list_level))
            inner = ''.join(inner_parts)
            return f'{indent}- {inner}'

        elif node_type == 'taskList':
            items = []
            for c in content:
                item_md = self._node_to_markdown(c, list_level)
                items.append(item_md)
            return '\n'.join(items)

        elif node_type == 'taskItem':
            indent = '  ' * list_level
            checked = attrs.get('checked', False)
            checkbox = '[x]' if checked else '[ ]'
            inner_parts = []
            for c in content:
                inner_parts.append(self._node_to_markdown(c, list_level))
            inner = ''.join(inner_parts)
            return f'{indent}- {checkbox} {inner}'

        elif node_type == 'blockquote':
            inner = '\n\n'.join(self._node_to_markdown(c) for c in content)
            lines = inner.split('\n')
            return '\n'.join(f'> {line}' for line in lines)

        elif node_type == 'codeBlock':
            language = attrs.get('language', '')
            code = self._render_inline_markdown(content)
            return f'```{language}\n{code}\n```'

        elif node_type == 'horizontalRule':
            return '---'

        elif node_type == 'image':
            src = attrs.get('src', '')
            alt = attrs.get('alt', '')
            return f'![{alt}]({src})'

        elif node_type == 'table':
            return self._table_to_markdown(content)

        elif node_type == 'text':
            text = node.get('text', '')
            marks = node.get('marks', [])
            for mark in marks:
                text = self._apply_mark_md(text, mark)
            return text

        elif node_type == 'hardBreak':
            return '  \n'

        elif node_type == 'mention':
            label = attrs.get('label', attrs.get('id', ''))
            return f'@{label}'

        else:
            if content:
                return '\n\n'.join(self._node_to_markdown(c) for c in content)
            return ''

    def _render_inline_markdown(self, content: List[Dict]) -> str:
        """Render inline content to Markdown."""
        parts = []
        for node in content:
            if node.get('type') == 'text':
                text = node.get('text', '')
                for mark in node.get('marks', []):
                    text = self._apply_mark_md(text, mark)
                parts.append(text)
            else:
                parts.append(self._node_to_markdown(node))
        return ''.join(parts)

    def _apply_mark_md(self, text: str, mark: Dict) -> str:
        """Apply a mark to text in Markdown."""
        mark_type = mark.get('type', '')
        attrs = mark.get('attrs', {})

        if mark_type == 'bold':
            return f'**{text}**'
        elif mark_type == 'italic':
            return f'*{text}*'
        elif mark_type == 'strike':
            return f'~~{text}~~'
        elif mark_type == 'code':
            return f'`{text}`'
        elif mark_type == 'link':
            href = attrs.get('href', '')
            return f'[{text}]({href})'
        else:
            return text

    def _table_to_markdown(self, rows: List[Dict]) -> str:
        """Convert table rows to Markdown table."""
        if not rows:
            return ''

        md_rows = []
        for i, row in enumerate(rows):
            if row.get('type') != 'tableRow':
                continue

            cells = row.get('content', [])
            cell_texts = []
            for cell in cells:
                cell_content = cell.get('content', [])
                text = ''
                for c in cell_content:
                    text += self._node_to_markdown(c)
                cell_texts.append(text.strip() or ' ')

            md_rows.append('| ' + ' | '.join(cell_texts) + ' |')

            # Add header separator after first row
            if i == 0:
                separator = '| ' + ' | '.join(['---'] * len(cell_texts)) + ' |'
                md_rows.append(separator)

        return '\n'.join(md_rows)

    # =========================================================================
    # EXPORT TO PDF
    # =========================================================================

    async def export_to_pdf(
        self,
        content: Dict[str, Any],
        title: str,
        include_header: bool = True,
        include_footer: bool = True,
        page_size: str = 'A4',
        margins: str = '2cm'
    ) -> bytes:
        """
        Export Tiptap content to PDF.

        Args:
            content: Tiptap JSON content
            title: Document title
            include_header: Add title at top
            include_footer: Add page numbers
            page_size: Paper size (A4, Letter)
            margins: Page margins

        Returns:
            PDF bytes
        """
        try:
            from weasyprint import HTML, CSS
        except ImportError:
            logger.error("WeasyPrint not installed. Install with: pip install weasyprint")
            raise ImportError("PDF export requires weasyprint package")

        html_content = self.tiptap_to_html(content)

        # Build full HTML document
        css = f"""
            @page {{
                size: {page_size};
                margin: {margins};
                @bottom-center {{
                    content: counter(page) " / " counter(pages);
                }}
            }}
            body {{
                font-family: 'Segoe UI', Arial, sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #333;
            }}
            h1 {{ font-size: 24pt; margin-bottom: 0.5em; }}
            h2 {{ font-size: 18pt; margin-bottom: 0.5em; }}
            h3 {{ font-size: 14pt; margin-bottom: 0.5em; }}
            p {{ margin-bottom: 0.8em; }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin: 1em 0;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #f5f5f5;
                font-weight: bold;
            }}
            code {{
                background-color: #f5f5f5;
                padding: 2px 4px;
                border-radius: 3px;
                font-family: 'Consolas', monospace;
                font-size: 10pt;
            }}
            pre {{
                background-color: #f5f5f5;
                padding: 12px;
                border-radius: 5px;
                overflow-x: auto;
            }}
            pre code {{
                padding: 0;
                background: none;
            }}
            blockquote {{
                border-left: 4px solid #ddd;
                margin: 1em 0;
                padding-left: 16px;
                color: #666;
            }}
            img {{
                max-width: 100%;
                height: auto;
            }}
            ul, ol {{
                margin-bottom: 1em;
                padding-left: 2em;
            }}
            .task-list {{
                list-style: none;
                padding-left: 0;
            }}
            .task-item {{
                margin-bottom: 0.5em;
            }}
            .mention {{
                background-color: #e3f2fd;
                padding: 2px 4px;
                border-radius: 3px;
            }}
            hr {{
                border: none;
                border-top: 1px solid #ddd;
                margin: 1.5em 0;
            }}
        """

        header_html = f'<h1 style="margin-bottom: 1em; border-bottom: 2px solid #333; padding-bottom: 0.5em;">{title}</h1>' if include_header else ''

        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{title}</title>
        </head>
        <body>
            {header_html}
            {html_content}
        </body>
        </html>
        """

        # Generate PDF
        html = HTML(string=full_html)
        pdf_bytes = html.write_pdf(stylesheets=[CSS(string=css)])

        return pdf_bytes

    # =========================================================================
    # EXPORT TO DOCX
    # =========================================================================

    async def export_to_docx(
        self,
        content: Dict[str, Any],
        title: str
    ) -> bytes:
        """
        Export Tiptap content to DOCX.

        Args:
            content: Tiptap JSON content
            title: Document title

        Returns:
            DOCX bytes
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise ImportError("DOCX export requires python-docx package")

        doc = Document()

        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Process content
        if content and 'content' in content:
            for node in content['content']:
                self._process_docx_node(doc, node)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.read()

    def _process_docx_node(self, doc, node: Dict, parent=None):
        """Process Tiptap node and add to DOCX document."""
        try:
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return

        node_type = node.get('type', '')
        attrs = node.get('attrs', {})
        content = node.get('content', [])

        if node_type == 'paragraph':
            para = doc.add_paragraph()
            self._add_runs_to_paragraph(para, content)

        elif node_type == 'heading':
            level = attrs.get('level', 1)
            heading = doc.add_heading(level=level)
            self._add_runs_to_paragraph(heading, content)

        elif node_type == 'bulletList':
            for item in content:
                if item.get('type') == 'listItem':
                    for c in item.get('content', []):
                        if c.get('type') == 'paragraph':
                            para = doc.add_paragraph(style='List Bullet')
                            self._add_runs_to_paragraph(para, c.get('content', []))

        elif node_type == 'orderedList':
            for item in content:
                if item.get('type') == 'listItem':
                    for c in item.get('content', []):
                        if c.get('type') == 'paragraph':
                            para = doc.add_paragraph(style='List Number')
                            self._add_runs_to_paragraph(para, c.get('content', []))

        elif node_type == 'blockquote':
            for c in content:
                para = doc.add_paragraph(style='Quote')
                if c.get('type') == 'paragraph':
                    self._add_runs_to_paragraph(para, c.get('content', []))

        elif node_type == 'codeBlock':
            para = doc.add_paragraph()
            code_text = ''.join(n.get('text', '') for n in content if n.get('type') == 'text')
            run = para.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)

        elif node_type == 'horizontalRule':
            para = doc.add_paragraph('_' * 50)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif node_type == 'table':
            self._add_docx_table(doc, content)

        elif node_type == 'image':
            src = attrs.get('src', '')
            # Note: Adding images from URLs would require downloading them first
            para = doc.add_paragraph(f'[Image: {src}]')

    def _add_runs_to_paragraph(self, para, content: List[Dict]):
        """Add text runs with formatting to a paragraph."""
        try:
            from docx.shared import Pt, RGBColor
        except ImportError:
            return

        for node in content:
            if node.get('type') == 'text':
                text = node.get('text', '')
                run = para.add_run(text)

                for mark in node.get('marks', []):
                    mark_type = mark.get('type', '')
                    if mark_type == 'bold':
                        run.bold = True
                    elif mark_type == 'italic':
                        run.italic = True
                    elif mark_type == 'underline':
                        run.underline = True
                    elif mark_type == 'strike':
                        run.font.strike = True
                    elif mark_type == 'code':
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)

            elif node.get('type') == 'hardBreak':
                para.add_run('\n')

    def _add_docx_table(self, doc, rows: List[Dict]):
        """Add a table to the DOCX document."""
        if not rows:
            return

        # Count columns from first row
        first_row = rows[0] if rows else {}
        num_cols = len(first_row.get('content', []))

        if num_cols == 0:
            return

        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'

        for i, row in enumerate(rows):
            if row.get('type') != 'tableRow':
                continue

            cells = row.get('content', [])
            for j, cell in enumerate(cells):
                if j < len(table.rows[i].cells):
                    cell_obj = table.rows[i].cells[j]
                    cell_content = cell.get('content', [])
                    text_parts = []
                    for c in cell_content:
                        if c.get('type') == 'paragraph':
                            for t in c.get('content', []):
                                if t.get('type') == 'text':
                                    text_parts.append(t.get('text', ''))
                    cell_obj.text = ''.join(text_parts)

    # =========================================================================
    # EXPORT TO HTML (STANDALONE)
    # =========================================================================

    async def export_to_html(
        self,
        content: Dict[str, Any],
        title: str,
        include_styles: bool = True
    ) -> str:
        """
        Export Tiptap content to standalone HTML.

        Args:
            content: Tiptap JSON content
            title: Document title
            include_styles: Include CSS styles

        Returns:
            HTML string
        """
        body_html = self.tiptap_to_html(content)

        styles = """
        <style>
            body {
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                max-width: 800px;
                margin: 40px auto;
                padding: 0 20px;
                line-height: 1.6;
                color: #333;
            }
            h1 { font-size: 2em; border-bottom: 2px solid #eee; padding-bottom: 0.3em; }
            h2 { font-size: 1.5em; }
            h3 { font-size: 1.25em; }
            code {
                background: #f5f5f5;
                padding: 2px 6px;
                border-radius: 3px;
                font-family: 'SF Mono', Consolas, monospace;
            }
            pre {
                background: #f5f5f5;
                padding: 16px;
                border-radius: 6px;
                overflow-x: auto;
            }
            pre code { padding: 0; background: none; }
            blockquote {
                border-left: 4px solid #ddd;
                margin: 1em 0;
                padding-left: 1em;
                color: #666;
            }
            table { border-collapse: collapse; width: 100%; margin: 1em 0; }
            th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            th { background: #f5f5f5; }
            img { max-width: 100%; }
            .task-list { list-style: none; padding-left: 0; }
            .mention { background: #e3f2fd; padding: 2px 4px; border-radius: 3px; }
        </style>
        """ if include_styles else ""

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    {styles}
</head>
<body>
    <h1>{title}</h1>
    {body_html}
</body>
</html>"""

    # =========================================================================
    # EXPORT TO MARKDOWN
    # =========================================================================

    async def export_to_markdown(
        self,
        content: Dict[str, Any],
        title: str,
        include_frontmatter: bool = False
    ) -> str:
        """
        Export Tiptap content to Markdown.

        Args:
            content: Tiptap JSON content
            title: Document title
            include_frontmatter: Add YAML frontmatter

        Returns:
            Markdown string
        """
        md_content = self.tiptap_to_markdown(content)

        if include_frontmatter:
            frontmatter = f"""---
title: {title}
date: {datetime.utcnow().isoformat()}
---

"""
            return frontmatter + f"# {title}\n\n" + md_content
        else:
            return f"# {title}\n\n{md_content}"

    # =========================================================================
    # EXPORT HISTORY
    # =========================================================================

    async def log_export(
        self,
        document_id: UUID,
        format: str,
        exported_by: UUID,
        file_path: Optional[str] = None,
        file_size: Optional[int] = None
    ) -> Dict[str, Any]:
        """Log an export operation."""
        conn = self._get_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        try:
            cur.execute("""
                INSERT INTO workspace.docs_exports (
                    document_id, format, file_path, file_size,
                    exported_by, exported_at
                ) VALUES (%s, %s, %s, %s, %s, NOW())
                RETURNING *
            """, (
                str(document_id), format, file_path, file_size, str(exported_by)
            ))

            row = cur.fetchone()
            conn.commit()

            return dict(row)

        except Exception as e:
            conn.rollback()
            logger.error(f"Failed to log export: {e}")
            raise
        finally:
            cur.close()
            conn.close()

    async def get_export_history(
        self,
        document_id: UUID,
        limit: int = 20
    ) -> List[Dict[str, Any]]:
        """Get export history for a document."""
        conn = self._get_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        try:
            cur.execute("""
                SELECT id, document_id, format, file_path, file_size,
                       exported_by, exported_at
                FROM workspace.docs_exports
                WHERE document_id = %s
                ORDER BY exported_at DESC
                LIMIT %s
            """, (str(document_id), limit))

            return [dict(row) for row in cur.fetchall()]

        finally:
            cur.close()
            conn.close()


# Singleton instance
_export_service: Optional[DocsExportService] = None


def get_docs_export_service() -> DocsExportService:
    """Get or create export service instance."""
    global _export_service
    if _export_service is None:
        _export_service = DocsExportService()
    return _export_service
