"""
Bheem OForms - OnlyOffice Document Forms Service

OnlyOffice-based fillable document forms service.
Creates DOCXF (form templates) and OFORM (fillable forms).

Features:
- DOCXF/OFORM file-based storage (Nextcloud)
- OnlyOffice Document Server for form creation and filling
- Version control
- Form responses collection
"""
from typing import Optional, Dict, Any, List
from uuid import UUID
from datetime import datetime, timedelta
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import text
import uuid
import jwt
import time
import hashlib
import io
import logging

from core.config import settings
from services.docs_storage_service import DocsStorageService, get_docs_storage_service

logger = logging.getLogger(__name__)


class OFormsService:
    """Service for OnlyOffice-based document forms."""

    def __init__(self, db: AsyncSession):
        self.db = db
        self.storage = get_docs_storage_service()

    async def create_oform(
        self,
        tenant_id: UUID,
        user_id: UUID,
        user_email: str = None,
        title: str = "Untitled Form",
        description: str = None,
        form_type: str = "docxf",  # docxf for template, oform for fillable
    ) -> Dict[str, Any]:
        """
        Create a new OnlyOffice document form.

        Args:
            tenant_id: Tenant UUID
            user_id: Creator's user UUID
            user_email: Creator's email for Nextcloud storage
            title: Form title
            description: Form description
            form_type: Form type (docxf for template, oform for fillable)

        Returns:
            Created form metadata
        """
        form_id = uuid.uuid4()
        document_key = f"oform-{form_id}-{int(time.time())}"

        # Get user email if not provided
        if not user_email:
            result = await self.db.execute(text("""
                SELECT email FROM workspace.tenant_users WHERE id = :user_id
            """), {"user_id": user_id})
            row = result.fetchone()
            user_email = row.email if row else f"user_{user_id}@bheem.cloud"

        # Create empty form file
        storage_path, file_size = await self._create_empty_form(
            form_id=form_id,
            tenant_id=tenant_id,
            user_email=user_email,
            title=title,
            form_type=form_type,
        )

        # Insert into database
        query = text("""
            INSERT INTO workspace.oforms (
                id, tenant_id, title, description, form_type,
                storage_path, storage_bucket, file_size,
                document_key, version,
                created_by, created_at, updated_at
            ) VALUES (
                :id, :tenant_id, :title, :description, :form_type,
                :storage_path, :storage_bucket, :file_size,
                :document_key, 1,
                :user_id, NOW(), NOW()
            )
            RETURNING id, created_at
        """)

        result = await self.db.execute(query, {
            "id": form_id,
            "tenant_id": tenant_id,
            "title": title,
            "description": description,
            "form_type": form_type,
            "storage_path": storage_path,
            "storage_bucket": "nextcloud",
            "file_size": file_size,
            "document_key": document_key,
            "user_id": user_id,
        })
        row = result.fetchone()
        await self.db.commit()

        # Record version history
        await self._record_version(
            form_id=form_id,
            version=1,
            user_id=user_id,
            storage_path=storage_path,
            file_size=file_size,
            comment="Initial version"
        )

        return {
            "id": str(form_id),
            "title": title,
            "description": description,
            "form_type": form_type,
            "storage_path": storage_path,
            "version": 1,
            "document_key": document_key,
            "file_size": file_size,
            "created_at": row.created_at.isoformat() if row.created_at else None,
            "edit_url": f"/oforms/{form_id}/edit",
        }

    async def _create_empty_form(
        self,
        form_id: UUID,
        tenant_id: UUID,
        user_email: str,
        title: str,
        form_type: str = "docxf",
    ) -> tuple:
        """Create an empty form template file."""
        from docx import Document

        # Create a new Word document as form template
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph('Add form fields using the OnlyOffice editor.')

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        file_content = buffer.read()
        file_size = len(file_content)

        # Use .docx extension for storage (Nextcloud compatibility)
        # OnlyOffice will handle form functionality based on fileType in config
        filename = f"{form_id}.docx"

        # Get Nextcloud username from email
        nc_username = self.storage.get_nextcloud_username(user_email)

        # Storage path in user's Forms folder (root level for permissions)
        folder_path = "Forms"
        storage_path = f"/{folder_path}/{form_id}.docx"

        # Upload to Nextcloud
        try:
            await self.storage.upload_file(
                file=io.BytesIO(file_content),
                filename=filename,
                tenant_id=tenant_id,
                folder_path=folder_path,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                nextcloud_user=nc_username,
            )
            logger.info(f"Created form template in Nextcloud: {storage_path} for user {nc_username}")
        except Exception as e:
            logger.error(f"Failed to upload form to Nextcloud: {e}")
            # Continue anyway - the form record will be created

        return storage_path, file_size

    async def _record_version(
        self,
        form_id: UUID,
        version: int,
        user_id: UUID,
        storage_path: str,
        file_size: int,
        comment: str = None,
    ):
        """Record a version in history."""
        try:
            await self.db.execute(text("""
                INSERT INTO workspace.oform_versions (
                    id, form_id, version, storage_path, file_size,
                    created_by, comment, created_at
                ) VALUES (
                    :id, :form_id, :version, :storage_path, :file_size,
                    :user_id, :comment, NOW()
                )
            """), {
                "id": uuid.uuid4(),
                "form_id": form_id,
                "version": version,
                "storage_path": storage_path,
                "file_size": file_size,
                "user_id": user_id,
                "comment": comment,
            })
            await self.db.commit()
        except Exception as e:
            logger.warning(f"Could not record version: {e}")

    async def get_oform(
        self,
        form_id: UUID,
        tenant_id: UUID,
        user_id: UUID,
    ) -> Optional[Dict[str, Any]]:
        """Get form metadata."""
        query = text("""
            SELECT
                f.id, f.title, f.description, f.form_type,
                f.storage_path, f.storage_bucket, f.file_size,
                f.version, f.document_key, f.status,
                f.is_starred, f.is_deleted, f.response_count,
                f.created_by, f.created_at, f.updated_at,
                tu.email as creator_email,
                tu.name as creator_name
            FROM workspace.oforms f
            LEFT JOIN workspace.tenant_users tu ON f.created_by = tu.id
            WHERE f.id = :form_id
            AND f.tenant_id = :tenant_id
            AND f.is_deleted = FALSE
        """)

        result = await self.db.execute(query, {
            "form_id": form_id,
            "tenant_id": tenant_id,
        })
        row = result.fetchone()

        if not row:
            return None

        return {
            "id": str(row.id),
            "title": row.title,
            "description": row.description,
            "form_type": row.form_type,
            "storage_path": row.storage_path,
            "storage_bucket": row.storage_bucket,
            "file_size": row.file_size,
            "version": row.version,
            "document_key": row.document_key,
            "status": row.status,
            "is_starred": row.is_starred,
            "response_count": row.response_count or 0,
            "created_by": str(row.created_by),
            "creator_email": row.creator_email,
            "creator_name": row.creator_name,
            "created_at": row.created_at.isoformat() if row.created_at else None,
            "updated_at": row.updated_at.isoformat() if row.updated_at else None,
        }

    async def list_oforms(
        self,
        tenant_id: UUID,
        user_id: UUID,
        search: Optional[str] = None,
        starred_only: bool = False,
        status: Optional[str] = None,
        limit: int = 50,
        offset: int = 0,
    ) -> Dict[str, Any]:
        """List forms with filtering."""
        conditions = ["f.tenant_id = :tenant_id", "f.is_deleted = FALSE"]
        params = {"tenant_id": tenant_id, "limit": limit, "offset": offset}

        if search:
            conditions.append("f.title ILIKE :search")
            params["search"] = f"%{search}%"

        if starred_only:
            conditions.append("f.is_starred = TRUE")

        if status:
            conditions.append("f.status = :status")
            params["status"] = status

        where_clause = " AND ".join(conditions)

        query = text(f"""
            SELECT
                f.id, f.title, f.description, f.form_type,
                f.version, f.file_size, f.status, f.response_count,
                f.is_starred, f.created_by, f.created_at, f.updated_at,
                tu.email as creator_email,
                tu.name as creator_name
            FROM workspace.oforms f
            LEFT JOIN workspace.tenant_users tu ON f.created_by = tu.id
            WHERE {where_clause}
            ORDER BY f.updated_at DESC
            LIMIT :limit OFFSET :offset
        """)

        result = await self.db.execute(query, params)
        rows = result.fetchall()

        # Get total count
        count_query = text(f"""
            SELECT COUNT(*) FROM workspace.oforms f
            WHERE {where_clause}
        """)
        count_result = await self.db.execute(count_query, params)
        total = count_result.scalar()

        forms = []
        for row in rows:
            forms.append({
                "id": str(row.id),
                "title": row.title,
                "description": row.description,
                "form_type": row.form_type,
                "version": row.version,
                "file_size": row.file_size,
                "status": row.status,
                "response_count": row.response_count or 0,
                "is_starred": row.is_starred,
                "created_by": str(row.created_by),
                "creator_email": row.creator_email,
                "creator_name": row.creator_name,
                "created_at": row.created_at.isoformat() if row.created_at else None,
                "updated_at": row.updated_at.isoformat() if row.updated_at else None,
            })

        return {
            "forms": forms,
            "total": total,
            "limit": limit,
            "offset": offset,
            "has_more": offset + len(forms) < total,
        }

    async def get_editor_config(
        self,
        form_id: UUID,
        tenant_id: UUID,
        user_id: UUID,
        user_name: str,
        user_email: str,
        mode: str = "edit",
    ) -> Dict[str, Any]:
        """
        Get OnlyOffice Document Editor configuration for forms.

        Args:
            form_id: Form UUID
            tenant_id: Tenant UUID for access validation
            user_id: User UUID
            user_name: Display name for collaboration
            user_email: User email
            mode: Edit mode (edit, view, fillForms)

        Returns:
            OnlyOffice editor configuration
        """
        # Get form metadata
        form = await self.get_oform(form_id, tenant_id, user_id)
        if not form:
            raise ValueError("Form not found")

        storage_path = form.get("storage_path")
        if not storage_path:
            raise ValueError("Form file not found")

        # Generate document key (unique per version for cache invalidation)
        document_key = f"{form_id}-v{form['version']}-{int(time.time())}"

        # Generate access token for document proxy endpoint
        access_token = jwt.encode(
            {
                "form_id": str(form_id),
                "exp": datetime.utcnow() + timedelta(hours=24),
            },
            settings.ONLYOFFICE_JWT_SECRET,
            algorithm="HS256"
        )

        if isinstance(access_token, bytes):
            access_token = access_token.decode('utf-8')

        from urllib.parse import quote
        encoded_token = quote(access_token, safe='')

        # Use proxy URL
        download_url = f"{settings.WORKSPACE_URL}/api/v1/oforms/{form_id}/content?token={encoded_token}"

        logger.info(f"Generated download URL for form {form_id}: {download_url[:100]}...")

        # Update document key in database
        await self.db.execute(
            text("UPDATE workspace.oforms SET document_key = :key WHERE id = :id"),
            {"key": document_key, "id": form_id}
        )
        await self.db.commit()

        # Callback URL for OnlyOffice
        callback_url = f"{settings.ONLYOFFICE_CALLBACK_URL.replace('/sheets', '/oforms')}/{form_id}/onlyoffice-callback"

        # Determine file type and document type
        form_type = form.get("form_type", "docxf")
        if form_type == "oform":
            file_type = "oform"
            document_type = "word"  # OFORM uses word document type
        else:
            file_type = "docxf"
            document_type = "word"  # DOCXF also uses word document type

        # Build editor config
        config = {
            "document": {
                "fileType": file_type,
                "key": document_key,
                "title": f"{form['title']}.{file_type}",
                "url": download_url,
                "permissions": {
                    "chat": True,
                    "comment": True,
                    "copy": True,
                    "download": True,
                    "edit": mode == "edit",
                    "fillForms": mode in ["edit", "fillForms"],
                    "print": True,
                    "review": mode in ["edit", "review"],
                },
            },
            "documentType": document_type,
            "editorConfig": {
                "callbackUrl": callback_url,
                "lang": "en",
                "mode": mode,
                "user": {
                    "id": str(user_id),
                    "name": user_name,
                },
                "customization": {
                    "autosave": True,
                    "comments": True,
                    "compactHeader": False,
                    "compactToolbar": False,
                    "feedback": False,
                    "forcesave": True,
                    "help": False,
                    "hideRightMenu": False,
                    "toolbarNoTabs": False,
                    "toolbarHideFileName": False,
                    "statusBar": True,
                    "leftMenu": True,
                    "rightMenu": True,
                    "toolbar": True,
                    "zoom": 100,
                    "logo": {
                        "image": f"{settings.WORKSPACE_URL}/static/bheem-logo.svg",
                        "imageDark": f"{settings.WORKSPACE_URL}/static/bheem-logo-dark.svg",
                        "imageLight": f"{settings.WORKSPACE_URL}/static/bheem-logo-light.svg",
                        "imageEmbedded": f"{settings.WORKSPACE_URL}/static/bheem-loader-logo.svg",
                        "url": f"{settings.WORKSPACE_URL}/oforms",
                        "visible": True
                    },
                    "customer": {
                        "address": "Bheem Cloud Services",
                        "info": "Bheem Workspace - Your productivity suite",
                        "logo": f"{settings.WORKSPACE_URL}/static/bheem-logo.svg",
                        "logoDark": f"{settings.WORKSPACE_URL}/static/bheem-logo-dark.svg",
                        "mail": "support@bheem.cloud",
                        "name": "Bheem Forms",
                        "www": "https://bheem.cloud"
                    },
                    "goback": {
                        "blank": False,
                        "text": "Back to Bheem",
                        "url": f"{settings.WORKSPACE_URL}/oforms"
                    },
                    "loaderLogo": f"{settings.WORKSPACE_URL}/static/bheem-loader-logo.svg",
                    "loaderName": "Bheem Forms",
                },
            },
            "height": "100%",
            "width": "100%",
            "type": "desktop",
        }

        # Add JWT token if enabled
        if settings.ONLYOFFICE_JWT_ENABLED and settings.ONLYOFFICE_JWT_SECRET:
            token = self._generate_onlyoffice_jwt(config)
            config["token"] = token

        return {
            "config": config,
            "documentServerUrl": settings.ONLYOFFICE_URL,
            "form": form,
        }

    def _generate_onlyoffice_jwt(self, payload: Dict[str, Any]) -> str:
        """Generate JWT token for OnlyOffice."""
        token = jwt.encode(
            payload,
            settings.ONLYOFFICE_JWT_SECRET,
            algorithm="HS256"
        )
        if isinstance(token, bytes):
            token = token.decode('utf-8')
        return token

    async def get_form_content(
        self,
        form_id: UUID,
        tenant_id: UUID,
    ) -> tuple:
        """Get form file content for OnlyOffice."""
        query = text("""
            SELECT f.storage_path, f.storage_bucket, f.form_type, f.title, tu.email as creator_email
            FROM workspace.oforms f
            LEFT JOIN workspace.tenant_users tu ON f.created_by = tu.id
            WHERE f.id = :form_id AND f.tenant_id = :tenant_id AND f.is_deleted = FALSE
        """)
        result = await self.db.execute(query, {"form_id": form_id, "tenant_id": tenant_id})
        row = result.fetchone()

        if not row:
            return None, None, None

        form_type = row.form_type or "docxf"
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # Try to get from storage first
        if row.storage_path:
            try:
                # Get Nextcloud username from creator's email
                nc_username = self.storage.get_nextcloud_username(row.creator_email) if row.creator_email else None

                # Download from Nextcloud
                file_obj, file_info = await self.storage.download_file(
                    storage_path=row.storage_path,
                    nextcloud_user=nc_username,
                )

                if file_obj:
                    # Read content from file object
                    content = file_obj.read() if hasattr(file_obj, 'read') else file_obj
                    return content, content_type, form_type
            except Exception as e:
                logger.warning(f"Could not get form from storage, generating template: {e}")

        # Fallback: Generate a fresh template document
        logger.info(f"Generating template for form {form_id}")
        try:
            from docx import Document

            doc = Document()
            doc.add_heading(row.title or "Untitled Form", 0)
            doc.add_paragraph('Add form fields using the OnlyOffice editor.')
            doc.add_paragraph('')
            doc.add_paragraph('Click on "Forms" tab in the toolbar to add form fields.')

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            content = buffer.read()

            return content, content_type, form_type
        except Exception as e:
            logger.error(f"Failed to generate form template: {e}")
            return None, None, None

    async def handle_onlyoffice_callback(
        self,
        form_id: UUID,
        callback_data: Dict[str, Any],
    ) -> Dict[str, int]:
        """Handle OnlyOffice Document Server callback."""
        status = callback_data.get("status")
        logger.info(f"OForms callback for {form_id}: status={status}")

        # Status codes:
        # 0 - no document with the key identifier
        # 1 - document is being edited
        # 2 - document is ready for saving
        # 3 - document saving error
        # 4 - document is closed with no changes
        # 6 - document is being edited, but the current document state is saved
        # 7 - error has occurred while force saving the document

        if status in [2, 6]:  # Ready for saving or force save
            download_url = callback_data.get("url")
            if download_url:
                await self._save_from_onlyoffice(form_id, download_url, callback_data)

        return {"error": 0}

    async def _save_from_onlyoffice(
        self,
        form_id: UUID,
        download_url: str,
        callback_data: Dict[str, Any],
    ):
        """Download and save form from OnlyOffice."""
        import aiohttp

        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(download_url, ssl=False) as response:
                    if response.status != 200:
                        logger.error(f"Failed to download form: {response.status}")
                        return

                    content = await response.read()

            # Get current form info
            result = await self.db.execute(text("""
                SELECT f.tenant_id, f.version, f.storage_path, f.form_type, tu.email as creator_email
                FROM workspace.oforms f
                LEFT JOIN workspace.tenant_users tu ON f.created_by = tu.id
                WHERE f.id = :form_id
            """), {"form_id": form_id})
            row = result.fetchone()

            if not row:
                logger.error(f"Form {form_id} not found")
                return

            # Get Nextcloud username
            nc_username = self.storage.get_nextcloud_username(row.creator_email) if row.creator_email else None

            # Upload to Nextcloud
            try:
                # Extract filename from storage path (use .docx for Nextcloud compatibility)
                filename = f"{form_id}.docx"

                await self.storage.upload_file(
                    file=io.BytesIO(content),
                    filename=filename,
                    tenant_id=row.tenant_id,
                    folder_path="Forms",
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    nextcloud_user=nc_username,
                )
                logger.info(f"Uploaded form {form_id} to Nextcloud")
            except Exception as e:
                logger.error(f"Failed to upload form to Nextcloud: {e}")

            # Update version
            new_version = row.version + 1
            new_key = f"oform-{form_id}-{int(time.time())}"

            await self.db.execute(text("""
                UPDATE workspace.oforms
                SET version = :version, document_key = :key, file_size = :size, updated_at = NOW()
                WHERE id = :form_id
            """), {
                "form_id": form_id,
                "version": new_version,
                "key": new_key,
                "size": len(content),
            })
            await self.db.commit()

            logger.info(f"Saved form {form_id} version {new_version}")

        except Exception as e:
            logger.error(f"Failed to save form from OnlyOffice: {e}")

    async def update_oform(
        self,
        form_id: UUID,
        tenant_id: UUID,
        user_id: UUID,
        title: Optional[str] = None,
        description: Optional[str] = None,
        status: Optional[str] = None,
    ) -> Optional[Dict[str, Any]]:
        """Update form metadata."""
        updates = []
        params = {"form_id": form_id, "tenant_id": tenant_id}

        if title is not None:
            updates.append("title = :title")
            params["title"] = title

        if description is not None:
            updates.append("description = :description")
            params["description"] = description

        if status is not None:
            updates.append("status = :status")
            params["status"] = status

        if not updates:
            return await self.get_oform(form_id, tenant_id, user_id)

        updates.append("updated_at = NOW()")
        update_clause = ", ".join(updates)

        await self.db.execute(text(f"""
            UPDATE workspace.oforms
            SET {update_clause}
            WHERE id = :form_id AND tenant_id = :tenant_id
        """), params)
        await self.db.commit()

        return await self.get_oform(form_id, tenant_id, user_id)

    async def delete_oform(
        self,
        form_id: UUID,
        tenant_id: UUID,
        user_id: UUID,
        permanent: bool = False,
    ) -> bool:
        """Delete a form (soft or permanent)."""
        if permanent:
            await self.db.execute(text("""
                DELETE FROM workspace.oforms
                WHERE id = :form_id AND tenant_id = :tenant_id
            """), {"form_id": form_id, "tenant_id": tenant_id})
        else:
            await self.db.execute(text("""
                UPDATE workspace.oforms
                SET is_deleted = TRUE, updated_at = NOW()
                WHERE id = :form_id AND tenant_id = :tenant_id
            """), {"form_id": form_id, "tenant_id": tenant_id})

        await self.db.commit()
        return True

    async def toggle_star(
        self,
        form_id: UUID,
        tenant_id: UUID,
        user_id: UUID,
    ) -> bool:
        """Toggle star status."""
        await self.db.execute(text("""
            UPDATE workspace.oforms
            SET is_starred = NOT is_starred, updated_at = NOW()
            WHERE id = :form_id AND tenant_id = :tenant_id
        """), {"form_id": form_id, "tenant_id": tenant_id})
        await self.db.commit()
        return True


def get_oforms_service(db: AsyncSession) -> OFormsService:
    """Factory function to get OForms service instance."""
    return OFormsService(db)
